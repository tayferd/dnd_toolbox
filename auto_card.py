import random
import docx
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Pt


import openai

openai.api_key = ""

num = input("How many characters?")
doc_name = input("Document name:")

races = ["Human", "Elf", "Dwarf", "Halfling", "Gnome", "Orc", "Tiefling", "Dragonborn", "Half-Elf", "Half-Orc", 
         "Aarakocra", "Genasi", "Goliath", "Aasimar", "Firbolg", "Kenku", "Tabaxi", "Triton", "Bugbear", "Goblin", 
         "Hobgoblin", "Kobold", "Yuan-Ti Pureblood", "Lizardfolk", "Tortle", "Changeling", "Kalashtar", "Shifter", 
         "Warforged", "Gith", "Centaur", "Loxodon", "Minotaur", "Simic Hybrid", "Vedalken", "Verdan", "Locathah", 
         "Leonin", "Satyr", "Abyssal Tiefling", "Dhampir", "Hexblood", "Reborn"]

melee_classes = ["Barbarian", "Bard", "Fighter", "Monk", "Paladin", "Ranger", "Rogue", "Cleric", "Artificer"]
spellcasting_classes = ["Bard", "Cleric", "Druid", "Sorcerer", "Warlock", "Paladin", "Ranger", "Wizard", "Artificer", "Blood Hunter"]
classes = ["Barbarian", "Bard", "Cleric", "Druid", "Fighter", "Monk", "Paladin", "Ranger", "Rogue", "Sorcerer", 
           "Warlock", "Wizard", "Artificer", "Blood Hunter"]

weapons = ["Axe", "Hammer", "Short Sword", "Bow", "Spear"]
armor_list = ["Chain Mail", "Leather", "Scalemail"]
fighter_abilities = ["Riposte", "Precision Attack", "Menacing Attack", "Trip Attack"]


# ... Rest of your code above ...

spells = {
    "Bard": {
        "Damage": ["Dissonant Whispers", "Vicious Mockery", "Thunderwave", "Heat Metal"],
        "Support": ["Hypnotic Pattern", "Invisibility", "Faerie Fire", "Heroism"],
        "Healing": ["Cure Wounds", "Healing Word", "Lesser Restoration", "Song of Rest"]
    },
    "Cleric": {
        "Damage": ["Inflict Wounds", "Sacred Flame", "Guiding Bolt", "Spiritual Weapon"],
        "Support": ["Bless", "Shield of Faith", "Protection from Evil and Good", "Sanctuary"],
        "Healing": ["Cure Wounds", "Healing Word", "Prayer of Healing", "Lesser Restoration"]
    },
    "Druid": {
        "Damage": ["Moonbeam", "Flame Blade", "Thorn Whip", "Call Lightning"],
        "Support": ["Barkskin", "Summon Fey", "Entangle", "Faerie Fire"],
        "Healing": ["Cure Wounds", "Healing Word", "Lesser Restoration", "Healing Spirit"]
    },
    "Paladin": {
        "Damage": ["Sacred Flame", "Bless", "Flame Blade", "Booming Blade"],
        "Support": ["Bless", "Shield of Faith", "Protection from Evil and Good", "Divine Favor"],
        "Healing": ["Cure Wounds", "Lay on Hands", "Lesser Restoration", "Aura of Vitality"]
    },
    "Ranger": {
        "Damage": ["Hail of Thorns", "Hunter's Mark", "Ensnaring Strike", "Conjure Barrage"],
        "Support": ["Barkskin", "Longstrider", "Pass without Trace", "Web"],
        "Healing": ["Cure Wounds", "Healing Spirit", "Lesser Restoration", "Goodberry"]
    },
    "Sorcerer": {
        "Damage": ["Chaos Bolt", "Chromatic Orb", "Burning Hands", "Lightning Bolt"],
        "Support": ["Wall of Fire", "Watery Sphere", "Control Person", "Haste", "Animate Objects", "Telekinesis", "Invisibility", "Polymorph"],
        "Healing": ["Cure Wounds", "Healing Word", "Lesser Restoration", "Revivify"]
    },
    "Warlock": {
        "Damage": ["Eldritch Blast", "Witch Bolt", "", "Arms of Hadar"],
        "Support": ["fear", "Hex", "Invisibility", "Fly"],
        "Healing": ["Cure Wounds", "Healing Word", "Lesser Restoration", "Whispers of the Grave"]
    },
    "Wizard": {
        "Damage": ["Magic Missile", "Fire Bolt", "Burning Hands", "Ray of Sickness"],
        "Support": ["Haste", "Shield", "Raise Undead", "telekenisis", "Invisibility", "Animate Objects", "Polymorph"],
        "Healing": [ "Lesser Restoration", "Falselife", "Vampiric Touch"]
    },
    "Artificer": {
        "Damage": ["Acid Arrow", "Explosive Cannon", "Magic Missle", "Ray of Enfeeblement"],
        "Support": ["Detect Magic", "Invisibility", "Arcane Lock", "Magic Mouth"],
        "Healing": ["Cure Wounds", "Healing Word", "Lesser Restoration", "Aura of Vitality"]
    },
    "Blood Hunter": {
        "Damage": ["Crimson Rite", "Blood Curse", "Vampiric Touch", "Wrathful Smite"],
        "Support": ["Blood Maledict", "Grim Psychometry", "Protection from Evil and Good", "Zone of Truth"],
        "Healing": ["Cure Wounds", "Healing Word", "Lesser Restoration", "Aura of Vitality"]
    },
}

rogue_abilities = [
    "Sneak Attack",
    "Cunning Action",
    "Uncanny Dodge",
    "Evasion",
    "Reliable Talent",
    "Blindsense",
    "Slippery Mind",
    "Elusive",
    "Stroke of Luck"
]
warlock_invocations = [
    "Agonizing Blast",
    "Eldritch Spear",
    "Devil's Sight",
    "Mask of Many Faces",
    "Thirsting Blade",
    "Repelling Blast",
    "Mire the Mind",
    "Bewitching Whispers",
    "Ascendant Step",
    "Lifedrinker",
    "Otherworldly Leap",
    "Whispers of the Grave"
    # Add more invocations as needed
]
monk_abilities = [
    "Unarmored Defense",
    "Martial Arts",
    "Ki",
    "Flurry of Blows",
    "Patient Defense",
    "Step of the Wind",
    "Deflect Missiles",
    "Slow Fall",
    "Stunning Strike",
    "Ki-Empowered Strikes",
    "Evasion",
    "Stillness of Mind",
    "Purity of Body",
    "Tongue of the Sun and Moon",
    "Empty Body",
    "Perfect Self"
    # Add more monk abilities as needed
]

sorcerer_metamagic = [
    "Careful Spell",
    "Distant Spell",
    "Empowered Spell",
    "Extended Spell",
    "Heightened Spell",
    "Quickened Spell",
    "Subtle Spell",
    "Twinned Spell"
]

blood_hunter_abilities = [
    "Crimson Rite",
    "Blood Maledict",
    "Hunter's Bane",
    "Brand of Castigation",
    "Sanguine Mastery",
    "Rite of the Dawn",
    "Brand of Tethering",
    "Rite of the Storm",
    "Brand of the Searing Strike",
    "Rite of the Frozen"
]

paladin_smites = [
    "Divine Smite",
    "Wrathful Smite",
    "Thunderous Smite",
    "Branding Smite",
    "Searing Smite",
    "Blinding Smite",
    "Staggering Smite",
    "Banishing Smite",
    "Crusader's Smite",
    "Eldritch Smite"
]

wizard_spells = [
    "Magic Missile",
    "Fire Bolt",
    "Burning Hands",
    "Ray of Sickness",
    "Thunderwave",
    "Chromatic Orb",
    "Feather Fall",
    "Mage Armor",
    "Shield",
    "Misty Step",
    "Invisibility",
    "Cure Wounds",
    "Healing Word",
    "Lesser Restoration",
    "Revivify",
    "Fly",
    "Counterspell",
    "Fireball",
    "Lightning Bolt",
    "Haste"
]

def get_abilities(char_class):
    if char_class == 'Fighter':
        return fighter_abilities
    elif char_class == 'Rogue':
        return random.sample(rogue_abilities, 3)
    elif char_class == 'Warlock':
        return random.sample(warlock_invocations, 4)
    elif char_class == 'Monk':
        return random.sample(monk_abilities, 4)
    elif char_class == 'Sorcerer':
        return random.sample(sorcerer_metamagic, 4)
    elif char_class == 'Blood Hunter':
        return random.sample(blood_hunter_abilities, 3)
    elif char_class == 'Paladin':
        return random.sample(paladin_smites, 4)
    elif char_class == 'Wizard':
        return random.sample(wizard_spells, 3)
    else:
        return []


def generate_character():
    race = random.choice(races)
    char_class = random.choice(classes)

    attack = random.choices(range(1, 9), k=1, weights=[3, 3, 3, 2, 2, 1, 1, 1])[0]
    defense = random.choices(range(1, 9), k=1, weights=[3, 3, 3, 2, 2, 1, 1, 1])[0]

    if char_class in melee_classes:
        weapon = random.choice(weapons)
        armor = random.choice(armor_list)
        abilities = get_abilities(char_class)
        damage_spell, support_spell, healing_spell = "Scroll", "Potion", "Medkit"
    elif char_class in spellcasting_classes:
        damage_spell = random.choice(spells[char_class]["Damage"])
        support_spell = random.choice(spells[char_class]["Support"])
        healing_spell = random.choice(spells[char_class]["Healing"])
        abilities = get_abilities(char_class)
        weapon, armor = "Dagger", "Robe",
    else:
        weapon, armor, damage_spell, support_spell, healing_spell, abilities = "N/A", "N/A", "N/A", "N/A", "N/A", []

    return race, char_class, attack, defense, weapon, armor, damage_spell, support_spell, healing_spell, abilities



def generate_prompt(race, char_class):
    return f"Generate a fantasy 5e name and special ability for a {race} {char_class}..."

# Create a new Word document
doc = docx.Document()

# Set font size and style
font_size = 12
font_style = "Arial"

# Add characters to the document with formatting
doc.add_heading('Characters', level=1)
for i in range(int(num)):
    # Generate a character
    race, char_class, attack, defense, weapon, armor, damage_spell, support_spell, healing_spell, abilities = generate_character()

    prompt = generate_prompt(race, char_class)

    # Get character details from OpenAI
    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        temperature=0.5,
        max_tokens=250
    )

    text = response.choices[0].text.strip()

    # Assuming the first line of the text is the character's name
    lines = text.split("\n")
    name = lines[0]
    story_text = "\n".join(lines[1:])

    # Create a new paragraph
    paragraph = doc.add_paragraph()

    # Set paragraph style
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Set font size and style
    run = paragraph.add_run(f"{name}\nRace: {race}\nClass: {char_class}\n\nMagic:\n"
                            f"Damage: {damage_spell if damage_spell else ''}\n"
                            f"Support: {support_spell if support_spell else 'Scrolls'}\n"
                            f"Healing: {healing_spell if healing_spell else 'Potions'}\n\n"
                            f"Weapon: {weapon if weapon else 'Dagger'}\n"
                            f"Armor: {armor if armor else 'Robes'}\n"
                            f"Abilities: {', '.join(abilities) if abilities else 'N/A'}\n\n"
                            f"{story_text}\n"
                            f"Attack: {attack}\n"
                            f"Defense: {defense}\n\n")
    font = run.font
    font.size = Pt(font_size)
    font.name = font_style

    # Add a page break after each character
    if i != int(num) - 1:
        doc.add_page_break()

# Save the document
doc.save(f'{doc_name}.docx')